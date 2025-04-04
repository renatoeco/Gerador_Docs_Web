import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
import pandas as pd
import os
import zipfile
import tempfile
from io import BytesIO
import shutil

from docx.oxml import OxmlElement
from docx.shared import RGBColor  # Para definir cor preta nas fontes
from docx.oxml.ns import qn

import streamlit.components.v1 as components

# HTML com a tag do GA4
ga4_tag = """
<!-- Google tag (gtag.js) -->
<script async src="https://www.googletagmanager.com/gtag/js?id=G-018YQ5WCV1"></script>
<script>
  window.dataLayer = window.dataLayer || [];
  function gtag(){dataLayer.push(arguments);}
  gtag('js', new Date());

  gtag('config', 'G-018YQ5WCV1');
</script>
"""

# Renderiza o HTML no frontend (dentro de um iframe invisível)
components.html(ga4_tag, height=0, width=0)



# Função do diálogo do resultado final
@st.dialog("Resultado")
def dialogo_resultado():
    """
    Chama a função para gerar os documentos em lote.
    """

    if st.session_state.success == "OK":
        sucesso()
    else:

        parametro1 = st.session_state.sheet_path
        parametro2 = st.session_state.model_path
        gerar_docs(parametro1, parametro2)





def adicionar_bordas_tabela(tabela):
    """Adiciona bordas de 1pt a todas as células de uma tabela."""
    for row in tabela.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')

            for border_tag in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_tag}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')  # 1 pt (8 em docx)
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)

            tcPr.append(tcBorders)


def gerar_docs(caminho_xlsx, caminho_docx):
    if "zip_buffer" in st.session_state:
        return

    temp_dir = tempfile.mkdtemp()
    doc_carregado = Document(caminho_docx)
    df_contratos = pd.read_excel(caminho_xlsx)

    # Cria a barra de progresso apenas uma vez
    progress_bar = st.progress(0, text="Iniciando...")

    for index, row in df_contratos.iterrows():
        
        # Atualiza a barra de progresso a cada iteração
        progress_bar.progress(index / len(df_contratos), text=f'Gerando documento {index + 1} de {len(df_contratos)}')
        
        st.session_state.cont += 1
        document = Document()
        document.styles['Normal'].font.name = 'Arial'

        # Cabeçalho com imagem
        cabecalho = document.sections[0].header
        paragrafo = cabecalho.paragraphs[0]
        if st.session_state.image_path != "no_image":
            paragrafo.alignment = 1
            paragrafo.add_run().add_picture(st.session_state.image_path, width=Inches(2))
        cabecalho.add_paragraph().space_after = Inches(0.5)

        # Rodapé com imagem
        footer = document.sections[0].footer
        paragrafo_footer = footer.paragraphs[0]
        if st.session_state.image_footer_path != "no_image":
            paragrafo_footer.alignment = 1
            paragrafo_footer.add_run().add_picture(st.session_state.image_footer_path, width=Inches(2))

        # Criar conteúdo na mesma ordem do modelo
        for element in doc_carregado.element.body:

            if element.tag.endswith('p'):
                paragrafo = next((p for p in doc_carregado.paragraphs if p._element is element), None)
                if paragrafo:
                    novo_paragrafo = document.add_paragraph()
                    novo_paragrafo.alignment = paragrafo.alignment

                    # Copiar estilo do parágrafo original
                    novo_paragrafo.style = paragrafo.style

                    for run in paragrafo.runs:
                        texto = run.text
                        for coluna in df_contratos.columns:
                            texto = texto.replace(f"{{{{{coluna}}}}}", str(row[coluna]))
                        novo_run = novo_paragrafo.add_run(texto)

                        # Copiar estilo do run original
                        novo_run.bold = run.bold
                        novo_run.italic = run.italic
                        novo_run.underline = run.underline
                        novo_run.font.size = run.font.size
                        novo_run.font.name = run.font.name
                        novo_run.font.color.rgb = run.font.color.rgb  # Copiar cor da fonte
                        novo_run.font.highlight_color = run.font.highlight_color  # Copiar cor de destaque

                    # Garantir que títulos, subtítulos e headings sejam pretos
                    if paragrafo.style.name.startswith("Heading") and novo_paragrafo.runs:
                        for novo_run in novo_paragrafo.runs:
                            novo_run.font.color.rgb = RGBColor(0, 0, 0)  # Preto

            elif element.tag.endswith('tbl'):
                tabela = next((t for t in doc_carregado.tables if t._element is element), None)
                if tabela:
                    nova_tabela = document.add_table(rows=len(tabela.rows), cols=len(tabela.columns))
                    nova_tabela.style = tabela.style

                    for i, row_tabela in enumerate(tabela.rows):
                        for j, cell in enumerate(row_tabela.cells):
                            texto_celula = cell.text
                            for coluna in df_contratos.columns:
                                texto_celula = texto_celula.replace(f"{{{{{coluna}}}}}", str(row[coluna]))
                            nova_tabela.cell(i, j).text = texto_celula

                            # Copiar cor da fonte da célula, se existir
                            if cell.paragraphs and cell.paragraphs[0].runs:
                                nova_tabela.cell(i, j).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Preto

                    # Aplicar bordas à nova tabela
                    adicionar_bordas_tabela(nova_tabela)

        # Nome do arquivo
        nome_arquivo = os.path.join(temp_dir, f"{st.session_state.docs_name}{row.iloc[0].replace('/', '_')}.docx")
        document.save(nome_arquivo)

    # Criar ZIP
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename in os.listdir(temp_dir):
            zip_file.write(os.path.join(temp_dir, filename), filename)

    zip_buffer.seek(0)
    st.session_state.zip_buffer = zip_buffer
    shutil.rmtree(temp_dir)

    st.balloons()
    st.session_state.success = "OK"
    sucesso()




def sucesso():

    if st.session_state.success == "OK":
        st.success(f'{st.session_state.cont} documentos gerados com sucesso!')

        st.download_button(
            label="Baixar ZIP com os arquivos .docx",
            data=st.session_state.zip_buffer,
            file_name="documentos_gerados.zip",
            mime="application/zip",
            type='primary'
        )


# Função principal
def main():


    # Inicializa a contagem de documentos
    if not "cont" in st.session_state: 
        st.session_state.cont = 0

    # Inicializa o marcador de sucesso
    if not "success" in st.session_state: 
        st.session_state.success = ""

    # Inicializa o prefixo do documento no session_state
    if not "docs_name" in st.session_state:
        st.session_state.docs_name = ""


    # Mostra o logotipo
    logo = "https://avatars.githubusercontent.com/u/85522293?v=4"
    
    with st.columns(3)[1]:
        st.image(logo, output_format="PNG", width=220)

    # Título e subtítulo
    st.markdown("<h1 style='text-align: center;'>Gerador de Documentos</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Para gerar documentos em lote, siga os passos a seguir:</p>", unsafe_allow_html=True)

    st.write('')
    st.write('')
    st.write('')

    
    # Passo 1 - Imagem do cabeçalho ----------------------------------------------
    st.markdown("<h5>1. Escolha uma imagem para o cabeçalho:</h5>", unsafe_allow_html=True)

    # Opção de seleção
    opcao_imagem = st.selectbox('', ('Continuar sem imagem no cabeçalho', 'Carregar imagem'))

    if opcao_imagem == "Carregar imagem":
        st.write('Lembre-se de ajustar o tamanho da imagem antes de carregar')
        arquivo_enviado = st.file_uploader("Escolha uma imagem (PNG, JPG ou JPEG)", type=["png", "jpg", "jpeg"])
        if arquivo_enviado:
            st.session_state.image_path = arquivo_enviado

            # Criando colunas para layout
            col1, col2 = st.columns([1, 3])

            col1.write("Imagem escolhida:")
            col2.image(st.session_state.image_path, width=150)


    elif opcao_imagem == "Continuar sem imagem no cabeçalho":
        st.session_state.image_path = "no_image"
    
    # Passo 2 - Carregar planilha ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>2. Carregue a tabela com as informações:</h5>", unsafe_allow_html=True)
    st.markdown("<span style='color:gray;'>Apenas são aceitas tabelas no formato .xlsx.</span>", unsafe_allow_html=True)
    st.markdown("<span style='color:gray;'>A <strong>primeira coluna da tabela</strong> será usada no <strong>nome dos arquivos</strong> gerados.</span>", unsafe_allow_html=True)

    # Upload do arquivo diretamente (sem botão separado)
    arquivo_enviado = st.file_uploader("Escolha a tabela (.xlsx)", type=["xlsx"])

    # Salva no session_state apenas se o usuário enviar um arquivo
    if arquivo_enviado is not None:
        st.session_state.sheet_path = arquivo_enviado

    # Passo 3 - Carregar modelo ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>3. Carregue o modelo do documento:</h5>", unsafe_allow_html=True)
    st.markdown("<span style='color:gray;'>Apenas são aceitos documentos no formato .docx.</span>", unsafe_allow_html=True)
    st.markdown("<span style='color:gray;'>As variárveis no modelo devem ter o mesmo nome das colunas da planilha, embaladas com chaves triplas {{{ }}}.</span>", unsafe_allow_html=True)

    # Upload do modelo diretamente (sem botão separado)
    arquivo_enviado = st.file_uploader("Escolha o modelo (.docx)", type=["docx"])

    # Salva no session_state apenas se o usuário enviar um arquivo
    if arquivo_enviado is not None:
        st.session_state.model_path = arquivo_enviado

    # Passo 4 - Imagem do rodapé ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>4. Escolha uma imagem para o rodapé:</h5>", unsafe_allow_html=True)
    
    opcao_imagem_rodape = st.selectbox('', ('Continuar sem imagem no rodapé', 'Carregar imagem'))

    if opcao_imagem_rodape == "Carregar imagem":
        st.write('Lembre-se de ajustar o tamanho da imagem antes de carregar')
        arquivo_enviado = st.file_uploader("Escolha uma imagem para o rodapé", type=["png", "jpg", "jpeg"])
        if arquivo_enviado:
            st.session_state.image_footer_path = arquivo_enviado

    elif opcao_imagem_rodape == "Continuar sem imagem no rodapé":
        st.session_state.image_footer_path = "no_image"

    # Verificando se as informações estão presentes no session_state
    if "image_footer_path" in st.session_state:
        if st.session_state.image_footer_path != "no_image":
            col1, col2 = st.columns([1, 3])
            col1.write("Imagem escolhida:")
            col2.image(st.session_state.image_footer_path, width=150)


    # Passo 5 - Escolha o nome dos arquivos ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>5. Deseja adicionar um prefixo ao nome dos documentos?</h5>", unsafe_allow_html=True)
    # if "docs_name" not in st.session_state:
    with st.form(key='prefixo_form', border=False):


        nome_doc_input = st.text_input('O nome dos documentos será: "prefixo" + "texto da primeira coluna da tabela". Exemplo: "Recibo_João Alves"')

        col1, col2 = st.columns([1, 3])

        submit_button = col1.form_submit_button(label='Confirmar prefixo')
        if submit_button and nome_doc_input != "":
            st.session_state.docs_name = nome_doc_input + "_"
            col2.write(':material/check:')



    # Passo 6 - Gerar documentos ----------------------------------------------
    st.markdown("<h5 style='padding-top: 40px'>6. Clique no botão abaixo para gerar os documentos:</h5>", unsafe_allow_html=True)

    if st.button("Gerar documentos!", type="primary"):
        if "sheet_path" in st.session_state and "model_path" in st.session_state:
          dialogo_resultado()
        else:
            st.warning('Você precisa escolher a TABELA DE INFORMAÇÕES e o DOCUMENTO MODELO (etapas 2 e 3).')   


if __name__ == "__main__":
    main()
